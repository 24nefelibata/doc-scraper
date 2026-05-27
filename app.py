# =============================================================================
# app.py — Documentation Scraper & Word Document Generator
# =============================================================================
# VERSION 3 CHANGES:
#   1. Full <table> support — reconstructed as real Word tables and aligned
#      ASCII tables in plain text export (critical for AI accuracy)
#   2. WebP image support — Pillow converts WebP → PNG before embedding
#   3. SVG images logged as [SVG DIAGRAM] placeholder instead of silently skipped
#   4. All v2 fixes retained (See Also stripping, unicode cleaning)
# =============================================================================

import streamlit as st
import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from PIL import Image as PilImage   # Pillow — for WebP → PNG conversion
import io
import re
import urllib3
import urllib.parse

# =============================================================================
# SECTION 1: PAGE CONFIG
# =============================================================================

st.set_page_config(
    page_title="Doc Scraper → Word / TXT",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="collapsed",
)

st.markdown("""
    <style>
        .main-title { font-size: 2.2rem; font-weight: 700; color: #1a1a2e; }
        .sub-title  { font-size: 1rem; color: #555; margin-top: -10px; }
        .stDownloadButton > button { font-weight: 600; border-radius: 8px; }
    </style>
""", unsafe_allow_html=True)

st.markdown('<p class="main-title">📄 Documentation Scraper → Word / TXT</p>',
            unsafe_allow_html=True)
st.markdown('<p class="sub-title">Paste documentation URLs below. One URL per line. '
            'Export as a formatted Word doc or clean plain text for AI ingestion.</p>',
            unsafe_allow_html=True)
st.divider()

urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

# =============================================================================
# SECTION 2: CONSTANTS & TEXT CLEANING
# =============================================================================

SEE_ALSO_PHRASES = [
    "see also", "related topics", "related links", "related articles",
    "further reading", "next steps", "what's next", "whats next",
    "additional resources", "learn more", "more information",
    "in this section", "on this page",
]

def clean_text(text: str) -> str:
    """
    Strips unicode artefacts and normalises whitespace.
    Handles the Â character common in GE Vernova docs (Latin-1/UTF-8 mismatch).
    """
    text = text.replace("\u00c2", "")    # Stray Â from encoding mismatch
    text = text.replace("\u00a0", " ")   # Non-breaking space → regular space
    text = text.replace("\u200b", "")    # Zero-width space
    text = text.replace("\u200c", "")    # Zero-width non-joiner
    text = text.replace("\u2019", "'")   # Right single quote → apostrophe
    text = text.replace("\u2018", "'")   # Left single quote → apostrophe
    text = text.replace("\u201c", '"')   # Left double quote
    text = text.replace("\u201d", '"')   # Right double quote
    text = re.sub(r" {2,}", " ", text)   # Collapse multiple spaces
    return text.strip()


def is_see_also_heading(text: str) -> bool:
    """Returns True if a heading signals the start of navigation-only content."""
    return any(phrase in text.lower() for phrase in SEE_ALSO_PHRASES)


# =============================================================================
# SECTION 3: TABLE EXTRACTION HELPER
# =============================================================================

def extract_table(table_tag) -> dict:
    """
    Converts a BeautifulSoup <table> tag into a structured dict.

    WHY A SEPARATE FUNCTION?
        Tables have their own internal structure (thead, tbody, tr, th, td)
        that doesn't fit the flat element list used for paragraphs and headings.
        We extract the full 2D grid here and let the Word/TXT builders
        render it appropriately for their format.

    RETURNS:
        {
          'type': 'table',
          'headers': ['Col1', 'Col2', ...],   # from <th> cells (may be empty list)
          'rows': [['val1', 'val2'], ...],     # from <td> cells
          'col_count': 4
        }

    HANDLES:
        - Tables with <thead>/<tbody> structure
        - Tables that use <th> inside <tbody> (first row as header)
        - Merged cells (colspan) — cell content is repeated for each spanned column
        - Nested content — extracts all text from inside each cell
    """

def is_layout_table(table_tag) -> bool:
    """
    Returns True if a <table> is being used for page layout rather than data.

    WHY THIS MATTERS:
        GE Vernova docs (and many older documentation sites) use <table> tags
        for two completely different purposes:
          1. DATA tables  — parameter lists, field definitions, comparison grids
          2. LAYOUT tables — positioning an image next to text, multi-column layouts

        Data tables should become Word tables.
        Layout tables should be transparently walked — extract their text and
        images as normal elements, ignoring the table structure entirely.

    SIGNALS THAT IT IS A LAYOUT TABLE:
        - Contains <img> tags directly inside cells
        - Has no <th> header cells anywhere
        - Has only 2-3 columns
        - Cells contain long paragraphs (over 200 chars) — data cells are short
        - Has a summary="" attribute (old HTML convention for layout tables)
        - Has role="presentation" attribute
    """
    # Explicit HTML attributes that declare layout intent
    if table_tag.get("role") == "presentation":
        return True
    if table_tag.get("summary") == "":
        return True

    # If any cell directly contains an <img>, it's a layout table
    for img in table_tag.find_all("img"):
        return True

    # If there are no <th> header cells at all, check cell content length
    has_th = bool(table_tag.find("th"))
    if not has_th:
        cells = table_tag.find_all("td")
        if cells:
            # Calculate average text length across all cells
            avg_len = sum(
                len(c.get_text(strip=True)) for c in cells
            ) / len(cells)
            # Data table cells are short (field names, values).
            # Layout table cells contain full paragraphs.
            if avg_len > 200:
                return True

    return False

    
    
    headers = []
    rows    = []

    # --- Extract header row from <thead> if it exists ---
    thead = table_tag.find("thead")
    if thead:
        for th in thead.find_all("th"):
            # colspan tells us if this header spans multiple columns
            colspan = int(th.get("colspan", 1))
            cell_text = clean_text(th.get_text(separator=" ", strip=True))
            # Repeat the header text for each column it spans
            headers.extend([cell_text] * colspan)

        # Some tables put <td> in thead instead of <th>
        if not headers:
            for td in thead.find_all("td"):
                colspan = int(td.get("colspan", 1))
                cell_text = clean_text(td.get_text(separator=" ", strip=True))
                headers.extend([cell_text] * colspan)

    # --- Extract data rows from <tbody> (or whole table if no tbody) ---
    tbody = table_tag.find("tbody") or table_tag

    for tr in tbody.find_all("tr", recursive=False):
        cells = tr.find_all(["td", "th"])
        if not cells:
            continue

        row = []
        for cell in cells:
            colspan   = int(cell.get("colspan", 1))
            cell_text = clean_text(cell.get_text(separator=" ", strip=True))
            row.extend([cell_text] * colspan)  # Expand spanned columns

        # If this row looks like a header row (all <th> cells) and we
        # haven't found headers yet, promote it to headers
        all_th = all(c.name == "th" for c in cells)
        if all_th and not headers:
            headers = row
        else:
            if any(cell for cell in row):  # Skip completely empty rows
                rows.append(row)

    # Determine column count from the widest row
    col_count = max(
        len(headers),
        max((len(r) for r in rows), default=0)
    )

    return {
        "type":      "table",
        "headers":   headers,
        "rows":      rows,
        "col_count": col_count,
    }


# =============================================================================
# SECTION 4: IMAGE FETCH WITH WEBP SUPPORT
# =============================================================================

BROWSER_HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
        "AppleWebKit/537.36 (KHTML, like Gecko) "
        "Chrome/124.0.0.0 Safari/537.36"
    )
}

def fetch_image(src_url: str) -> tuple:
    """
    Downloads an image and returns (bytes, format_hint).
    Converts WebP → PNG using Pillow so python-docx can embed it.
    Returns (None, None) on any failure.

    WHY PILLOW FOR WEBP?
        python-docx uses the underlying python-docx image handler which
        does not support WebP format. Pillow opens ANY image format it
        knows (WebP, TIFF, BMP, etc.) and can re-save it as PNG in memory,
        which python-docx handles perfectly.

    RETURNS:
        (image_bytes: bytes, hint: str)  where hint is 'png', 'jpeg', or 'skip'
        (None, None) if download or conversion failed
    """
    # Skip SVG — it's a vector format, not a raster image.
    # python-docx cannot embed SVGs at all.
    if src_url.lower().endswith(".svg") or "svg" in src_url.lower().split("?")[0]:
        return (None, "svg")

    try:
        resp = requests.get(src_url, headers=BROWSER_HEADERS,
                            timeout=10, verify=False)
        resp.raise_for_status()

        content_type = resp.headers.get("Content-Type", "").lower()

        # If it's not actually an image (e.g., the server returned HTML
        # for a broken URL), skip it
        if "image" not in content_type and not src_url.lower().split("?")[0].endswith(
            (".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".tiff")
        ):
            return (None, None)

        raw_bytes = resp.content

        # --- WEBP / UNKNOWN FORMAT: run through Pillow → PNG ---
        # We do this for WebP explicitly, and also as a fallback for any
        # format python-docx might not recognise.
        if "webp" in content_type or src_url.lower().endswith(".webp"):
            try:
                img = PilImage.open(io.BytesIO(raw_bytes)).convert("RGBA")
                out = io.BytesIO()
                img.save(out, format="PNG")
                out.seek(0)
                return (out.read(), "png")
            except Exception:
                return (None, None)

        # --- STANDARD FORMATS (JPEG, PNG, GIF) ---
        # Attempt a quick Pillow open to verify the bytes are a valid image
        try:
            PilImage.open(io.BytesIO(raw_bytes)).verify()
        except Exception:
            # Corrupt or unreadable image — skip silently
            return (None, None)

        return (raw_bytes, "ok")

    except Exception:
        return (None, None)


# =============================================================================
# SECTION 5: CORE SCRAPING FUNCTION
# =============================================================================

def scrape_page(url: str) -> dict:
    """
    Fetches a URL and returns a structured list of content elements.

    Element types returned:
        h1–h6   : heading with 'text'
        p       : paragraph with 'text'
        li      : list item with 'text' and 'list_type' ('ul'/'ol')
        code    : code block with 'text'
        blockquote : indented quote with 'text'
        table   : full table dict from extract_table()
        image   : image with 'src' and 'alt'
        svg     : placeholder for SVG diagrams
    """
    try:
        response = requests.get(url, headers=BROWSER_HEADERS,
                                timeout=20, verify=False)
        response.raise_for_status()
    except requests.exceptions.RequestException as e:
        return {"url": url, "error": str(e), "elements": []}

    soup = BeautifulSoup(response.text, "html.parser")

    title_tag  = soup.find("title")
    page_title = clean_text(title_tag.get_text(strip=True)) if title_tag else url

    # --- Strip noise elements ---
    noise_selectors = [
        "nav", "footer", "header", "aside",
        "script", "style",
        "[class*='sidebar']",      "[class*='nav']",
        "[class*='menu']",         "[class*='toc']",
        "[class*='cookie']",       "[class*='banner']",
        "[id*='sidebar']",         "[id*='nav']",
        "[class*='seealso']",      "[class*='see-also']",
        "[class*='related']",      "[class*='relatedLinks']",
        "[class*='relatedTopics']","[class*='footer-links']",
        "[id*='seealso']",         "[id*='related']",
    ]
    for selector in noise_selectors:
        for tag in soup.select(selector):
            tag.decompose()

    # --- Find main content area ---
    content_area = (
        soup.find("main")
        or soup.find("article")
        or soup.find(class_=lambda c: c and any(
            kw in c for kw in ["content", "article", "post-body",
                                "docs-content", "markdown-body", "entry-content"]
        ))
        or soup.find("body")
    )

    if content_area is None:
        return {"url": url, "title": page_title,
                "error": "Could not find content area.", "elements": []}

    elements        = []
    stop_extraction = False  # Flipped True when a "See Also" heading is hit

    def walk(node):
        nonlocal stop_extraction

        for child in node.children:
            if stop_extraction:
                return

            if not hasattr(child, "name") or child.name is None:
                continue

            tag_name = child.name.lower()

            # ---- HEADINGS ----
            if tag_name in ("h1", "h2", "h3", "h4", "h5", "h6"):
                text = clean_text(child.get_text(separator=" ", strip=True))
                if is_see_also_heading(text):
                    stop_extraction = True
                    return
                if text:
                    elements.append({"type": tag_name, "text": text})

           # ---- TABLES ----
            elif tag_name == "table":
                if is_layout_table(child):
                    # Layout table — ignore the table structure entirely.
                    # Recurse into it so images and text inside cells
                    # are extracted as normal individual elements.
                    walk(child)
                else:
                    # Data table — convert to a structured Word table.
                    table_data = extract_table(child)
                    if table_data["rows"] or table_data["headers"]:
                        elements.append(table_data)

            # ---- PARAGRAPHS ----
            elif tag_name == "p":
                text = clean_text(child.get_text(separator=" ", strip=True))
                if text and len(text) > 2:
                    elements.append({"type": "p", "text": text})

            # ---- CODE BLOCKS ----
            elif tag_name in ("pre", "code"):
                text = clean_text(child.get_text(strip=True))
                if text:
                    elements.append({"type": "code", "text": text})

            # ---- BLOCKQUOTES ----
            elif tag_name == "blockquote":
                text = clean_text(child.get_text(separator=" ", strip=True))
                if text:
                    elements.append({"type": "blockquote", "text": text})

            # ---- UNORDERED LISTS ----
            elif tag_name == "ul":
                for li in child.find_all("li", recursive=False):
                    text = clean_text(li.get_text(separator=" ", strip=True))
                    if text:
                        elements.append({"type": "li", "text": text,
                                         "list_type": "ul"})

            # ---- ORDERED LISTS ----
            elif tag_name == "ol":
                for li in child.find_all("li", recursive=False):
                    text = clean_text(li.get_text(separator=" ", strip=True))
                    if text:
                        elements.append({"type": "li", "text": text,
                                         "list_type": "ol"})

            # ---- IMAGES ----
            elif tag_name == "img":
                src = child.get("src", "")
                alt = clean_text(child.get("alt", "Image"))
                if src:
                    abs_src = urllib.parse.urljoin(url, src)
                    # Flag SVGs separately so the builder can log a placeholder
                    if abs_src.lower().endswith(".svg"):
                        elements.append({"type": "svg", "alt": alt})
                    else:
                        elements.append({"type": "image",
                                         "src": abs_src, "alt": alt})

            # ---- RECURSE INTO CONTAINERS ----
            elif tag_name in ("div", "section", "main", "article", "figure",
                              "details", "summary"):
                walk(child)
            # Note: we no longer recurse into td/th — extract_table() handles those

    walk(content_area)

    return {"url": url, "title": page_title, "elements": elements}


# =============================================================================
# SECTION 6: WORD TABLE BUILDER HELPER
# =============================================================================

def add_word_table(doc, table_data: dict):
    """
    Renders a table element dict as a proper python-docx Table.

    HOW python-docx TABLES WORK:
        doc.add_table(rows=N, cols=M) creates a grid.
        table.cell(row_index, col_index) gives you a cell.
        cell.text = "value" sets the cell content.
        We then apply a built-in Word table style and shade the header row.

    COLUMN WIDTH:
        We divide the available page width (5.6 inches after margins) equally
        among all columns. This keeps the table within the page margins.
    """
    headers   = table_data["headers"]
    rows      = table_data["rows"]
    col_count = table_data["col_count"]

    if col_count == 0:
        return

    has_header_row = bool(headers)
    total_rows     = (1 if has_header_row else 0) + len(rows)

    if total_rows == 0:
        return

    # Create the table grid
    table = doc.add_table(rows=total_rows, cols=col_count)

    # Apply a clean built-in Word table style
    # "Table Grid" gives solid borders on all cells — clean and readable
    table.style = "Table Grid"

    # --- Fill header row ---
    if has_header_row:
        header_row = table.rows[0]
        for col_idx, header_text in enumerate(headers[:col_count]):
            cell = header_row.cells[col_idx]
            cell.text = header_text

            # Make header text bold
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)

            # Shade header cells with the same blue used in the GE Vernova table
            # We do this by injecting the XML shading element directly
            tc   = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd  = OxmlElement("w:shd")
            shd.set(qn("w:val"),   "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"),  "D9E1F2")  # Light blue — matches GE docs style
            tcPr.append(shd)

    # --- Fill data rows ---
    row_offset = 1 if has_header_row else 0
    for row_idx, row_data in enumerate(rows):
        word_row = table.rows[row_idx + row_offset]
        for col_idx, cell_text in enumerate(row_data[:col_count]):
            cell = word_row.cells[col_idx]
            cell.text = cell_text
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    # Add a blank paragraph after the table for visual spacing
    doc.add_paragraph()


# =============================================================================
# SECTION 7: PLAIN TEXT TABLE RENDERER
# =============================================================================

def render_text_table(table_data: dict) -> str:
    """
    Renders a table as an ASCII grid for the plain text export.

    WHY ASCII TABLES FOR AI?
        An AI reading plain text needs structure to understand relationships
        between cells. A properly aligned ASCII table like:

            | Order | Field Name | Required | Values         |
            |-------|------------|----------|----------------|
            | 0     | Result Set | X        | 5              |
            | 1     | Prod Unit  | X        | PU_Id          |

        is far more parseable than:
            Order Field Name Required Values
            0 Result Set X 5
            1 Prod Unit X PU_Id

        The pipe characters and dashes give the AI clear column boundaries,
        which dramatically improves its ability to answer structured questions
        like "which fields are required?" or "what does field 3 reference?"
    """
    headers   = table_data["headers"]
    rows      = table_data["rows"]
    col_count = table_data["col_count"]

    if col_count == 0:
        return ""

    # Combine headers + rows into one list for width calculation
    all_rows = []
    if headers:
        all_rows.append(headers)
    all_rows.extend(rows)

    # Pad every row to col_count columns (handle missing cells gracefully)
    all_rows = [row + [""] * (col_count - len(row)) for row in all_rows]

    # Calculate the maximum width needed for each column
    col_widths = []
    for col_idx in range(col_count):
        max_width = max(
            len(str(row[col_idx])) for row in all_rows
        ) if all_rows else 5
        col_widths.append(max(max_width, 3))  # Minimum 3 chars wide

    def format_row(row):
        """Formats one row as | cell | cell | cell |"""
        cells = []
        for col_idx, cell in enumerate(row[:col_count]):
            # Left-align, pad to column width
            cells.append(str(cell).ljust(col_widths[col_idx]))
        return "| " + " | ".join(cells) + " |"

    def separator_row():
        """Creates the |---|---|---| divider line"""
        return "|-" + "-|-".join("-" * w for w in col_widths) + "-|"

    lines = []

    if headers:
        lines.append(format_row(headers))
        lines.append(separator_row())
        for row in rows:
            padded = row + [""] * (col_count - len(row))
            lines.append(format_row(padded))
    else:
        # No header — just rows with a separator after the first
        for i, row in enumerate(rows):
            padded = row + [""] * (col_count - len(row))
            lines.append(format_row(padded))
            if i == 0:
                lines.append(separator_row())

    return "\n".join(lines)


# =============================================================================
# SECTION 8: WORD DOCUMENT BUILDER
# =============================================================================

def add_hyperlink(paragraph, text: str, url: str):
    """Inserts a clickable hyperlink into a python-docx paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def build_word_document(pages: list, doc_title: str) -> io.BytesIO:
    """
    Assembles all scraped pages into a formatted Word .docx file.
    v3: adds real table rendering and WebP image support.
    """
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)

    # Cover title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run(doc_title)
    run.bold = True
    run.font.size = Pt(26)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    title_para.space_after = Pt(6)

    sub = doc.add_paragraph("Compiled Documentation")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(12)
    sub.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()

    for i, page in enumerate(pages):
        if i > 0:
            doc.add_page_break()

        page_heading = doc.add_heading(page.get("title", page["url"]), level=1)
        if page_heading.runs:
            page_heading.runs[0].font.color.rgb = RGBColor(0x00, 0x72, 0xC6)

        source_para = doc.add_paragraph("Source: ")
        source_para.runs[0].font.size = Pt(9)
        source_para.runs[0].italic = True
        source_para.runs[0].font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        add_hyperlink(source_para, page["url"], page["url"])
        doc.add_paragraph()

        if "error" in page and not page.get("elements"):
            err_para = doc.add_paragraph(f"⚠ Could not retrieve content: {page['error']}")
            if err_para.runs:
                err_para.runs[0].font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            continue

        for elem in page.get("elements", []):
            etype = elem["type"]

            # ---- HEADINGS ----
            if etype in ("h1", "h2", "h3", "h4", "h5", "h6"):
                level_num  = int(etype[1])
                word_level = min(level_num + 1, 9)
                doc.add_heading(elem["text"], level=word_level)

            # ---- PARAGRAPHS ----
            elif etype == "p":
                doc.add_paragraph(elem["text"])

            # ---- TABLES (NEW IN v3) ----
            elif etype == "table":
                add_word_table(doc, elem)

            # ---- CODE BLOCKS ----
            elif etype == "code":
                code_para = doc.add_paragraph(elem["text"])
                code_run  = (code_para.runs[0] if code_para.runs
                             else code_para.add_run(elem["text"]))
                code_run.font.name = "Courier New"
                code_run.font.size = Pt(9)

            # ---- BLOCKQUOTES ----
            elif etype == "blockquote":
                bq_para = doc.add_paragraph(elem["text"])
                bq_para.paragraph_format.left_indent = Inches(0.5)
                if bq_para.runs:
                    bq_para.runs[0].italic = True

            # ---- LIST ITEMS ----
            elif etype == "li":
                style = ("List Number" if elem.get("list_type") == "ol"
                         else "List Bullet")
                doc.add_paragraph(elem["text"], style=style)

            # ---- IMAGES (v3: WebP support via Pillow) ----
            elif etype == "image":
                img_bytes, hint = fetch_image(elem["src"])
                if img_bytes:
                    img_stream = io.BytesIO(img_bytes)
                    try:
                        doc.add_picture(img_stream, width=Inches(5))
                        if elem.get("alt") and elem["alt"] != "Image":
                            cap = doc.add_paragraph(elem["alt"])
                            if cap.runs:
                                cap.runs[0].italic = True
                                cap.runs[0].font.size = Pt(9)
                                cap.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
                    except Exception:
                        pass

            # ---- SVG PLACEHOLDER (NEW IN v3) ----
            elif etype == "svg":
                alt_text = elem.get("alt", "Diagram")
                svg_para = doc.add_paragraph(f"[SVG Diagram: {alt_text}]")
                if svg_para.runs:
                    svg_para.runs[0].italic = True
                    svg_para.runs[0].font.color.rgb = RGBColor(0x88, 0x88, 0x88)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# =============================================================================
# SECTION 9: PLAIN TEXT BUILDER
# =============================================================================

def build_plain_text(pages: list, doc_title: str) -> io.BytesIO:
    """
    Converts scraped pages into clean plain text optimised for AI input.
    v3: tables rendered as ASCII grids; SVG logged as placeholder.
    """
    lines = []
    lines.append("=" * 60)
    lines.append(doc_title.upper())
    lines.append("=" * 60)
    lines.append("")

    for page in pages:
        lines.append("=" * 60)
        lines.append(page.get("title", page["url"]))
        lines.append(f"Source: {page['url']}")
        lines.append("=" * 60)
        lines.append("")

        if "error" in page and not page.get("elements"):
            lines.append(f"[ERROR: {page['error']}]")
            lines.append("")
            continue

        for elem in page.get("elements", []):
            etype = elem["type"]

            if etype == "h1":
                lines.append(f"# {elem['text']}")
                lines.append("")
            elif etype == "h2":
                lines.append(f"## {elem['text']}")
                lines.append("")
            elif etype == "h3":
                lines.append(f"### {elem['text']}")
                lines.append("")
            elif etype in ("h4", "h5", "h6"):
                lines.append(f"#### {elem['text']}")
                lines.append("")
            elif etype == "p":
                lines.append(elem["text"])
                lines.append("")
            elif etype == "li":
                lines.append(f"  • {elem['text']}")
            elif etype == "code":
                lines.append("[CODE]")
                lines.append(elem["text"])
                lines.append("[/CODE]")
                lines.append("")
            elif etype == "blockquote":
                lines.append(f'  "{elem["text"]}"')
                lines.append("")

            # ---- TABLE: render as ASCII grid (NEW IN v3) ----
            elif etype == "table":
                ascii_table = render_text_table(elem)
                if ascii_table:
                    lines.append(ascii_table)
                    lines.append("")

            # ---- IMAGE: alt text note ----
            elif etype == "image":
                if elem.get("alt") and elem["alt"] != "Image":
                    lines.append(f"[IMAGE: {elem['alt']}]")

            # ---- SVG: placeholder (NEW IN v3) ----
            elif etype == "svg":
                lines.append(f"[SVG DIAGRAM: {elem.get('alt', 'see source page')}]")

        lines.append("")

    full_text = "\n".join(lines)
    buffer    = io.BytesIO(full_text.encode("utf-8"))
    buffer.seek(0)
    return buffer


# =============================================================================
# SECTION 10: STREAMLIT USER INTERFACE
# =============================================================================

col_left, col_right = st.columns([2, 1])

with col_left:
    st.subheader("① Enter URLs")
    url_input = st.text_area(
        label="One URL per line:",
        height=220,
        placeholder=(
            "https://docs.example.com/getting-started\n"
            "https://docs.example.com/installation\n"
            "https://docs.example.com/api-reference"
        ),
        help="Paste as many URLs as you like. They appear in the document in this order.",
    )

    doc_title_input = st.text_input(
        label="Document Title (appears on the first page)",
        value="Documentation Compilation",
    )

    btn_col1, btn_col2 = st.columns(2)
    with btn_col1:
        word_button = st.button(
            "📄 Export as Word (.docx)",
            type="primary",
            use_container_width=True,
            help="Formatted document with real tables, headings, bullets, and images.",
        )
    with btn_col2:
        txt_button = st.button(
            "🤖 Export as Plain Text (.txt)",
            type="secondary",
            use_container_width=True,
            help="Clean text with ASCII tables. Optimised for AI tools.",
        )

with col_right:
    st.subheader("ℹ How it works")
    st.info(
        "**📄 Word export** — real tables with shaded headers, "
        "embedded images (PNG/JPEG/WebP), headings, bullets.\n\n"
        "**🤖 Plain Text export** — tables as aligned ASCII grids "
        "(pipe-separated columns), images as alt-text notes. "
        "Ideal for ChatGPT, Claude, or Gemini.\n\n"
        "---\n"
        "**Filtered out automatically:**\n"
        "- Navigation menus & sidebars\n"
        "- 'See Also' / 'Related Topics' sections\n"
        "- Cookie banners & footers\n"
        "- Unicode artefacts (stray Â characters)\n"
        "- SVG images (shown as placeholder label)"
    )

st.divider()

# =============================================================================
# SECTION 11: SHARED SCRAPING + EXPORT LOGIC
# =============================================================================

compile_triggered = word_button or txt_button
export_mode       = "word" if word_button else "txt"

if compile_triggered:

    if not url_input.strip():
        st.error("Please enter at least one URL before exporting.")
        st.stop()

    raw_urls   = [line.strip() for line in url_input.strip().splitlines()]
    valid_urls = [u for u in raw_urls if u.startswith("http")]

    if not valid_urls:
        st.error("No valid URLs found. Each URL must start with http:// or https://")
        st.stop()

    skipped = [u for u in raw_urls if u and not u.startswith("http")]
    if skipped:
        st.warning(f"Skipped {len(skipped)} line(s) that didn't look like URLs.")

    mode_label = "Word document" if export_mode == "word" else "plain text file"
    st.info(f"Found **{len(valid_urls)} URL(s)**. Building {mode_label}…")

    scraped_pages = []
    progress_bar  = st.progress(0, text="Initialising…")

    with st.status("Scraping pages…", expanded=True) as status_box:
        for idx, url in enumerate(valid_urls):
            fraction = idx / len(valid_urls)
            progress_bar.progress(
                fraction,
                text=f"Scraping {idx + 1} of {len(valid_urls)}: {url}"
            )
            st.write(f"🔍 Fetching: `{url}`")
            result = scrape_page(url)
            scraped_pages.append(result)

            if "error" in result and not result.get("elements"):
                st.write(f"  ⚠ Error: {result['error']}")
            else:
                elem_count   = len(result.get("elements", []))
                table_count  = sum(1 for e in result.get("elements", [])
                                   if e["type"] == "table")
                image_count  = sum(1 for e in result.get("elements", [])
                                   if e["type"] == "image")
                st.write(
                    f"  ✅ Done — {elem_count} elements "
                    f"({table_count} table(s), {image_count} image(s))"
                )

        progress_bar.progress(1.0, text="Scraping complete. Building output…")
        status_box.update(label="Scraping complete!", state="complete",
                          expanded=False)

    safe_filename = (
        doc_title_input.strip().replace(" ", "_").replace("/", "-")
        or "documentation"
    )

    # --- WORD EXPORT ---
    if export_mode == "word":
        with st.spinner("Assembling Word document (fetching images…)"):
            docx_buffer = build_word_document(scraped_pages, doc_title_input)

        st.success(f"✅ Word document compiled from {len(valid_urls)} page(s)!")
        st.download_button(
            label="⬇ Download Word Document (.docx)",
            data=docx_buffer,
            file_name=f"{safe_filename}.docx",
            mime=(
                "application/vnd.openxmlformats-officedocument"
                ".wordprocessingml.document"
            ),
            use_container_width=True,
        )

    # --- PLAIN TEXT EXPORT ---
    elif export_mode == "txt":
        with st.spinner("Assembling plain text file…"):
            txt_buffer = build_plain_text(scraped_pages, doc_title_input)

        st.success(f"✅ Plain text file compiled from {len(valid_urls)} page(s)!")

        txt_buffer.seek(0)
        preview_text  = txt_buffer.read().decode("utf-8")
        preview_lines = "\n".join(preview_text.splitlines()[:60])

        with st.expander("👁 Preview first 60 lines"):
            st.code(preview_lines, language=None)

        txt_buffer.seek(0)
        st.download_button(
            label="⬇ Download Plain Text (.txt)",
            data=txt_buffer,
            file_name=f"{safe_filename}.txt",
            mime="text/plain",
            use_container_width=True,
        )
        st.info(
            "💡 **AI tip:** Open the .txt file, press `Ctrl+A` → `Ctrl+C`, "
            "then paste it into your AI chat with your question. "
            "The pipe-separated tables give the AI clear column boundaries "
            "so it can accurately answer questions about parameters and values."
        )
